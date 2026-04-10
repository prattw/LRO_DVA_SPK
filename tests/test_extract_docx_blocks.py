"""DOCX block classification (uses minimal python-docx document)."""

from __future__ import annotations

from io import BytesIO

import docx

from vantage_preprocess.extract.docx_extract import classify_paragraph_kind, extract_docx_blocks
from vantage_preprocess.extract.schemas import DocxBlockKind


def test_extract_docx_preserves_heading_order() -> None:
    d = docx.Document()
    d.add_heading("Title", level=1)
    d.add_paragraph("Body one.")
    d.add_heading("Section", level=2)
    d.add_paragraph("Body two.")
    buf = BytesIO()
    d.save(buf)
    content = buf.getvalue()

    result = extract_docx_blocks(content)
    assert len(result.blocks) == 4
    assert result.blocks[0].kind == DocxBlockKind.HEADING
    assert result.blocks[0].text == "Title"
    assert result.blocks[1].kind == DocxBlockKind.PARAGRAPH
    assert result.blocks[1].text == "Body one."
    assert result.blocks[2].kind == DocxBlockKind.HEADING
    assert result.blocks[3].text == "Body two."
    flat = result.flattened_text()
    assert "Title" in flat and "Body two" in flat


def test_classify_paragraph_kind_body() -> None:
    d = docx.Document()
    p = d.add_paragraph("x")
    assert classify_paragraph_kind(p) == DocxBlockKind.PARAGRAPH
