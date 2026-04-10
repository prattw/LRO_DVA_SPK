"""Section / heading detection (DOCX styles + PDF heuristics)."""

from __future__ import annotations

from io import BytesIO

import docx

from vantage_preprocess.extract.engine import extract_docx_document
from vantage_preprocess.models.document import PageText, StructuredDocument
from vantage_preprocess.models.enums import DocumentType, ExtractMethod
from vantage_preprocess.sections.detect import detect_sections
from vantage_preprocess.sections.models import HeadingSource
from vantage_preprocess.sections.patterns import score_line_as_heading
from vantage_preprocess.services.sectionize import section_blocks_from_document


def test_score_csi_and_part_patterns() -> None:
    s1 = score_line_as_heading("SECTION 09 65 00 RESILIENT FLOORING")
    assert s1 is not None and s1.score >= 5.0
    s2 = score_line_as_heading("PART 1 - GENERAL")
    assert s2 is not None and s2.score >= 5.0
    s3 = score_line_as_heading("PART 2 - PRODUCTS")
    assert s3 is not None
    s4 = score_line_as_heading("1.1 SUMMARY")
    assert s4 is not None and s4.score >= 3.0


def test_docx_styles_populated_on_extract() -> None:
    d = docx.Document()
    d.add_heading("Scope", level=2)
    d.add_paragraph("Body under scope.")
    d.add_heading("Products", level=2)
    d.add_paragraph("Product text.")
    buf = BytesIO()
    d.save(buf)
    doc = extract_docx_document(buf.getvalue(), "spec.docx", "/tmp/spec.docx")
    assert doc.docx_extraction is not None
    det = detect_sections(doc)
    assert det.strategy == "docx_styles"
    assert len(det.sections) >= 2
    titles = [s.heading_text for s in det.sections if s.heading_text]
    assert "Scope" in titles and "Products" in titles
    assert all(s.heading_source == HeadingSource.DOCX_STYLE for s in det.sections if s.heading_text)


def test_pdf_multipage_headings_preserve_pages() -> None:
    doc = StructuredDocument(
        document_id="a" * 32,
        source_filename="x.pdf",
        source_path="/x",
        source_sha256="b" * 64,
        mime_type="application/pdf",
        document_type=DocumentType.PDF,
        pages=[
            PageText(
                page_number=1,
                text="PART 1 - GENERAL\nIntro text here.",
                method=ExtractMethod.PARSE,
                confidence=1.0,
            ),
            PageText(
                page_number=2,
                text="PART 2 - PRODUCTS\nProduct requirements.",
                method=ExtractMethod.PARSE,
                confidence=1.0,
            ),
        ],
        overall_extract_method=ExtractMethod.PARSE,
        overall_confidence=1.0,
    )
    det = detect_sections(doc)
    assert det.strategy == "text_heuristic"
    assert len(det.sections) >= 2
    pages_spanned = {(s.page_start, s.page_end) for s in det.sections}
    assert (1, 1) in pages_spanned or any(s.page_start <= 1 <= s.page_end for s in det.sections)
    assert any(s.page_end >= 2 for s in det.sections)


def test_heading_noise_falls_back_to_per_page() -> None:
    lines = [f"{i}.1 SECTION TITLE LINE" for i in range(60)]
    text = "\n".join(lines)
    doc = StructuredDocument(
        document_id="a" * 32,
        source_filename="noisy.pdf",
        source_path="/n",
        source_sha256="b" * 64,
        mime_type="application/pdf",
        document_type=DocumentType.PDF,
        pages=[
            PageText(
                page_number=1,
                text=text[: len(text) // 2],
                method=ExtractMethod.PARSE,
                confidence=1.0,
            ),
            PageText(
                page_number=2,
                text=text[len(text) // 2 :],
                method=ExtractMethod.PARSE,
                confidence=1.0,
            ),
        ],
        overall_extract_method=ExtractMethod.PARSE,
        overall_confidence=1.0,
    )
    det = detect_sections(doc)
    assert det.strategy == "fallback_page"
    assert len(det.sections) == 2
    assert all(s.heading_source == HeadingSource.FALLBACK_PAGE for s in det.sections)


def test_section_blocks_from_document_matches_detect() -> None:
    doc = StructuredDocument(
        document_id="a" * 32,
        source_filename="t.txt",
        source_path="/t",
        source_sha256="b" * 64,
        mime_type="text/plain",
        document_type=DocumentType.TXT,
        pages=[
            PageText(
                page_number=1,
                text="DIVISION 03\nConcrete work paragraph.",
                method=ExtractMethod.PARSE,
                confidence=1.0,
            ),
        ],
        overall_extract_method=ExtractMethod.PARSE,
        overall_confidence=1.0,
    )
    blocks = section_blocks_from_document(doc)
    assert len(blocks) >= 1
    assert blocks[0].text
    assert blocks[0].heading_confidence is not None
